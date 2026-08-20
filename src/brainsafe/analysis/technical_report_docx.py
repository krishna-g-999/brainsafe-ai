"""Convert the technical report to Word, with the flowcharts rendered rather than pasted as code.

Pandoc does not know what a mermaid block is. Converted naively, every flowchart in the report
arrives in Word as a grey box of `flowchart TD` source, which is worse than omitting it: it looks
like a rendering failure and it tells the reader nothing.

So each mermaid block is rendered to PNG by the mermaid CLI first and replaced by an image
reference, and pandoc then embeds real pictures. The CLI is fetched through npx, which needs node
and a network on the first run only; if it is unavailable the block is replaced by a readable
plain-text rendering of the same steps rather than by source code, and the run says which.

Figures already generated for the manuscript are appended as an annex, so the Word document carries
the statistical evidence alongside the prose rather than referring to files the reader does not
have.

Output: docs/TECHNICAL_REPORT.docx

Run:  python src/brainsafe/analysis/technical_report_docx.py
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[3]
SRC = ROOT / "docs" / "TECHNICAL_REPORT.md"
OUT = ROOT / "docs" / "TECHNICAL_REPORT.docx"
FIGDIR = ROOT / "manuscript" / "figures"

# Figures appended as an annex, with the caption each carries in the manuscript.
ANNEX = [
    ("Figure10_endpoint_selection.png",
     "Endpoint selection. (A) Data volume against cross-validated discrimination across the panel. "
     "(B) The condition that actually decides deployment: whether a threshold exists that recovers "
     "actives without firing on unrelated chemistry. Endpoints that rank well can still be "
     "withheld. (C) The same data bar applied to every candidate target surveyed."),
    ("Figure1_architecture.png",
     "How a query is answered. Every target score is admitted only in proportion to the predicted "
     "probability that the compound reaches the brain, and surviving scores are ranked by "
     "enrichment over each endpoint's base rate rather than by raw probability."),
    ("Figure9_model_atlas.png",
     "The panel, one mark per estimator, so that no claim rests on a mean a reader cannot check. "
     "Marker shape carries the metric: AUROC and R-squared both run to 1.0 and are not the same "
     "quantity."),
    ("Figure6_validation.png",
     "Four validations that a cross-validated score cannot replace: calibration, recall on withheld "
     "scaffold classes, specificity on chemistry the server should stay quiet about, and the "
     "adversarial suite including the check that fails."),
    ("Figure7_binder_panel.png",
     "Every binder endpoint named, deployed and withdrawn alike, with what it discriminates and "
     "what it recovers."),
    ("Figure5_negative_class.png",
     "Recovery of the measured negative class from censored bounds, and its effect per endpoint."),
]


def render_mermaid(code: str, dest: Path) -> bool:
    """Render one mermaid block to PNG. False if the renderer is unavailable."""
    with tempfile.TemporaryDirectory() as tmp:
        mmd = Path(tmp) / "d.mmd"
        mmd.write_text(code, encoding="utf-8")
        cmd = ["npx", "-y", "@mermaid-js/mermaid-cli@11", "-i", str(mmd), "-o", str(dest),
               "-b", "white", "-s", "2"]
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=300, shell=True)
        except Exception:
            return False
        return r.returncode == 0 and dest.exists() and dest.stat().st_size > 0


def as_text(code: str) -> str:
    """A readable fallback: the flow as indented steps, not as source."""
    lines = []
    for ln in code.splitlines():
        ln = ln.strip()
        if not ln or ln.startswith(("flowchart", "graph", "%%")):
            continue
        ln = re.sub(r"<br\s*/?>", " ", ln)
        m = re.match(r'^(\w+)\s*(?:\[|\{\{|\{|\()([^\]\}\)]*)', ln)
        arrow = re.findall(r'-->\|([^|]*)\||-->', ln)
        label = m.group(2).strip('"') if m else ln
        label = re.sub(r'[\[\]\{\}\(\)"]', "", label).strip()
        if label:
            cond = f" [{arrow[0]}]" if arrow and arrow[0] else ""
            lines.append(f"    - {label}{cond}")
    return "\n".join(lines)


def polish(path: Path) -> None:
    """Apply what a stylesheet cannot: header-row repetition, header shading, column widths.

    These are properties of each table instance rather than of the style, so they cannot be set in
    the reference document and have to be applied after pandoc has written the file.
    """
    import docx
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    d = docx.Document(str(path))
    for tbl in d.tables:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        # Let Word size the columns to their content instead of the fixed grid pandoc emits, which
        # is what makes one long text column squeeze every other column to nothing.
        tblpr = tbl._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "autofit")
        tblpr.append(layout)

        head = tbl.rows[0]
        trpr = head._tr.get_or_add_trPr()
        rep = OxmlElement("w:tblHeader")            # repeat on every page the table spans
        rep.set(qn("w:val"), "true")
        trpr.append(rep)
        for cell in head.cells:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "EAF0F7")
            cell._tc.get_or_add_tcPr().append(shd)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(8.5)
        # A very wide table is set a point smaller so it still fits the text column.
        if len(tbl.columns) >= 6:
            for row in tbl.rows[1:]:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(7.8)
    d.save(str(path))


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"{SRC} not found; run build_technical_report.py first")
    text = SRC.read_text(encoding="utf-8")

    imgdir = ROOT / "docs" / "_report_figures"
    imgdir.mkdir(parents=True, exist_ok=True)
    blocks = list(re.finditer(r"```mermaid\n(.*?)\n```", text, re.S))
    print(f"{len(blocks)} flowchart(s) to render")

    rendered = failed = 0
    for i, m in enumerate(reversed(blocks), 1):
        idx = len(blocks) - i + 1
        dest = imgdir / f"flowchart_{idx}.png"
        if render_mermaid(m.group(1), dest):
            rel = dest.relative_to(ROOT / "docs").as_posix()
            # Size by aspect. Pandoc caps an image at the text width and keeps its aspect, so a
            # bare height attribute is ignored and a vertical chain still runs to 40 cm across
            # three pages. The width that yields the wanted height is computed and given instead.
            from PIL import Image
            with Image.open(dest) as im:
                w, h = im.size
            MAX_H_CM, TEXT_W_CM = 20.5, 14.8
            want_w = MAX_H_CM * w / max(h, 1)
            attrs = ("{width=%.1fcm}" % want_w) if want_w < TEXT_W_CM else "{width=100%}"
            repl = f"![Flowchart {idx}]({rel}){attrs}"
            rendered += 1
            print(f"  rendered flowchart {idx} -> {dest.name} ({w}x{h}) {attrs}")
        else:
            repl = ("**Flowchart " + str(idx) + "** (renderer unavailable; steps in order)\n\n"
                    + as_text(m.group(1)))
            failed += 1
            print(f"  flowchart {idx}: renderer unavailable, written as text")
        text = text[:m.start()] + repl + text[m.end():]

    annex = ["\n\n---\n\n## Annex: figures\n"]
    for i, (name, cap) in enumerate(ANNEX, 1):
        p = FIGDIR / name
        if not p.exists():
            print(f"  annex figure missing, skipped: {name}")
            continue
        annex.append(f"\n**Figure A{i}.** {cap}\n\n"
                     f"![Figure A{i}]({p.as_posix()}){{width=100%}}\n")
    text += "\n".join(annex)

    tmp_md = ROOT / "docs" / "_technical_report_for_docx.md"
    tmp_md.write_text(text, encoding="utf-8")
    try:
        import pypandoc
        args = ["--resource-path", str(ROOT / "docs"), "--resource-path", str(ROOT),
                "--toc", "--toc-depth=3"]
        ref = ROOT / "docs" / "report_reference.docx"
        if not ref.exists():
            import report_reference_docx
            report_reference_docx.main()
        args += ["--reference-doc", str(ref)]
        pypandoc.convert_file(str(tmp_md), "docx", outputfile=str(OUT), extra_args=args)
    finally:
        tmp_md.unlink(missing_ok=True)

    polish(OUT)

    size = OUT.stat().st_size / 1e6
    print(f"\nwrote {OUT.relative_to(ROOT)}  ({size:.2f} MB)")
    print(f"  flowcharts rendered as images : {rendered}")
    if failed:
        print(f"  flowcharts written as text    : {failed}")
    print(f"  annex figures embedded        : {sum(1 for n, _ in ANNEX if (FIGDIR / n).exists())}")


if __name__ == "__main__":
    main()

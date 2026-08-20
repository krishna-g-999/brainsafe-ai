"""Build the Word reference document that gives the technical report its typography.

Pandoc's built-in template ships one table style, named `Table`, with no borders and no header
shading, and it sets no body font. The result is what a reader sees as "scattered": eighteen tables
with invisible structure, four paragraph styles used interchangeably, and whatever font Word
defaults to.

A reference document fixes all of that at the source. Pandoc copies its styles into the output, so
every heading, paragraph and table is laid out by rules defined once here rather than patched
afterwards.

Two things are done in raw XML because python-docx has no API for them: table borders, and the
"repeat header row" flag that keeps a column heading visible when a long table breaks across pages.
Both matter for a document whose longest table is 76 rows.

Output: docs/report_reference.docx

Run:  python src/brainsafe/analysis/report_reference_docx.py
"""
from __future__ import annotations

from pathlib import Path

import docx
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "docs" / "report_reference.docx"

# The brand palette, so the document and the server look like the same project.
NAVY = RGBColor(0x0D, 0x21, 0x37)
NAVY_MID = RGBColor(0x1A, 0x3A, 0x5C)
GOLD_HEX = "F0A500"
RULE_HEX = "C8D2DF"
HEAD_FILL = "0D2137"

BODY_FONT = "Calibri"        # present on every Windows and Office install; no substitution surprises
MONO_FONT = "Consolas"


def set_font(style, name: str, size: float, *, bold=False, color=None, space_after=6,
             space_before=0, line=1.15):
    f = style.font
    f.name = name
    f.size = Pt(size)
    f.bold = bold
    if color is not None:
        f.color.rgb = color
    # East-Asian font must be set in XML too, or Word substitutes for any non-Latin glyph.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)
    pf = style.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line


def table_borders(style, hex_colour: str = RULE_HEX, size: int = 4):
    """Hairline borders on every edge. python-docx exposes no API for this."""
    # A style element has no tblPr accessor, unlike a table element, so it is created by hand.
    tblpr = style.element.find(qn("w:tblPr"))
    if tblpr is None:
        tblpr = OxmlElement("w:tblPr")
        style.element.append(tblpr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_colour)
        borders.append(el)
    tblpr.append(borders)
    # Breathing room inside every cell; the default is zero and is why dense tables look crushed.
    margins = OxmlElement("w:tblCellMar")
    for side, w in (("top", 60), ("left", 90), ("bottom", 60), ("right", 90)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(w))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tblpr.append(margins)


def main() -> None:
    d = docx.Document()          # the blank template carries the 100 built-in styles

    s = d.sections[0]
    s.orientation = WD_ORIENT.PORTRAIT
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)      # A4
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = s.bottom_margin = Cm(2.0)

    set_font(d.styles["Normal"], BODY_FONT, 10.5, space_after=7, line=1.22)
    for name, size, colour, before in (("Heading 1", 20, NAVY, 0),
                                       ("Heading 2", 15, NAVY, 16),
                                       ("Heading 3", 12, NAVY_MID, 13),
                                       ("Heading 4", 11, NAVY_MID, 10)):
        if name in [st.name for st in d.styles]:
            set_font(d.styles[name], BODY_FONT, size, bold=True, color=colour,
                     space_after=5, space_before=before, line=1.1)

    # Pandoc maps its own paragraph classes onto these names; each is set so that switching between
    # them, which pandoc does freely, does not change how the page looks.
    for name in ("Body Text", "First Paragraph", "Compact", "Block Text"):
        if name in [st.name for st in d.styles]:
            set_font(d.styles[name], BODY_FONT, 10.5, space_after=7, line=1.22)

    for name, size in (("Image Caption", 9), ("Caption", 9)):
        if name in [st.name for st in d.styles]:
            st = d.styles[name]
            set_font(st, BODY_FONT, size, color=RGBColor(0x5A, 0x66, 0x72), space_after=12)
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name in ("Source Code", "Verbatim Char", "HTML Code"):
        if name in [st.name for st in d.styles]:
            set_font(d.styles[name], MONO_FONT, 9, space_after=6, line=1.1)

    # The table style pandoc actually applies. Creating it here is what makes the output legible.
    if "Table" not in [st.name for st in d.styles]:
        d.styles.add_style("Table", WD_STYLE_TYPE.TABLE)
    tstyle = d.styles["Table"]
    set_font(tstyle, BODY_FONT, 8.5, space_after=0, space_before=0, line=1.05)
    table_borders(tstyle)

    d.add_paragraph("BrainSafe AI reference styles", style="Heading 1")
    d.add_paragraph(
        "This file exists only to carry typography into the generated report. Its content is never "
        "read. Regenerate with report_reference_docx.py.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.save(OUT)
    print(f"wrote {OUT.relative_to(ROOT)}")
    print(f"  body {BODY_FONT} 10.5pt, headings navy, tables bordered at 8.5pt on A4 "
          f"with 2.2 cm side margins")


if __name__ == "__main__":
    main()

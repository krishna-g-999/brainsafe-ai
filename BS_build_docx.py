# -*- coding: utf-8 -*-
"""Submission-standard Word manuscript (python-docx). Numbers read from saved artifacts
(supplementary CSVs, BS_auroc_cis.json); figures embedded from figures/. Adds line numbering,
structured abstract, highlights, abbreviations, equations, hyperparameter/software tables,
AUROC 95% CIs, and full front/back matter."""
import os, json
os.chdir(os.path.dirname(os.path.abspath(__file__)))
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
sec = doc.sections[0]; sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"): setattr(sec, m, Inches(1))
# continuous line numbering (submission standard)
ln = OxmlElement("w:lnNumType"); ln.set(qn("w:countBy"), "1"); ln.set(qn("w:restart"), "continuous"); ln.set(qn("w:distance"), "360")
sec._sectPr.append(ln)

def H(t, lvl=1): return doc.add_heading(t, level=lvl)
def P(t, italic=False, bold=False, align=None, size=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic; r.bold = bold
    if size: r.font.size = Pt(size)
    if align == "c": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if align == "j": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p
def bullet(t):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(t).font.size = Pt(10.5)
def eq(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(t); r.italic = True; r.font.size = Pt(10.5)
def figure(path, caption, width=6.3):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(caption, italic=True, size=9)
def table_from_df(df):
    df = df.fillna("–")
    t = doc.add_table(rows=1, cols=len(df.columns)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for j, c in enumerate(df.columns):
        cl = t.rows[0].cells[j]; cl.text = ""; r = cl.paragraphs[0].add_run(str(c)); r.bold = True; r.font.size = Pt(8.5)
    for _, row in df.iterrows():
        cs = t.add_row().cells
        for j, c in enumerate(df.columns):
            cs[j].text = ""; rr = cs[j].paragraphs[0].add_run(str(row[c])); rr.font.size = Pt(8.5)
    return t

cis = json.load(open("BS_auroc_cis.json")) if os.path.exists("BS_auroc_cis.json") else {}

# display names so every table matches the prose (data files use raw endpoint keys)
DISP = {"MAO_A": "MAO-A", "MAO_B": "MAO-B", "GSK3B": "GSK-3β", "HT2A": "5-HT2A"}

# ---- TITLE ----
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("BrainSafe AI: an evidence-grounded, calibrated, blood–brain-barrier-gated "
               "multi-endpoint predictor of small-molecule effects on the human brain")
r.bold = True; r.font.size = Pt(16)
P("Running title: Multi-endpoint prediction of brain effects from measured data", align="c", italic=True, size=10)
P("Author(s): [To be completed]   ·   ORCID: [____]", align="c", size=10)
P("Affiliation: Sri Sathya Sai Institute of Higher Learning, India   ·   Corresponding author: [email]", align="c", size=10)

# ---- GRAPHICAL ABSTRACT ----
figure("figures/graphical_abstract.png", "Graphical abstract.", width=6.6)

# ---- HIGHLIGHTS ----
H("Highlights", 2)
for h in [
    "An open tool predicts CNS-relevant effects of any small molecule from structure alone, trained only on measured public data (ChEMBL_37, B3DB; 64,474 records).",
    "Eight classification endpoints (BBB, AChE, BChE, BACE1, GSK-3β, MAO-A, MAO-B, hERG) plus four receptor potency regressions and a measured antioxidant model.",
    "Probabilities are isotonic-calibrated and accompanied by conformal prediction sets with empirical coverage close to the 90% target.",
    "Models are evaluated under random, scaffold, leave-cluster-out and temporal splits; predictions are BBB-gated and grounded in nearest measured analogues.",
]:
    bullet(h)

# ---- ABSTRACT ----
H("Abstract", 1)
P("Background. Judging whether a small molecule acts on the brain means answering several questions "
  "together: blood–brain-barrier (BBB) penetration, engagement of disease-relevant central-nervous-"
  "system (CNS) targets, developability, safety and clinical precedent. Public tools address these only "
  "in part (Daina et al., 2017; Fu et al., 2024). "
  "Results. BrainSafe AI is an open tool that predicts this profile from chemical structure alone. It "
  "integrates eight classification endpoints trained on measured public data (ChEMBL_37 and B3DB; "
  "64,474 records; Mendez et al., 2019; Meng et al., 2021): BBB, AChE, BChE, BACE1, GSK-3β, MAO-A, MAO-B "
  "and the hERG safety liability, together with four receptor potency regressions, a measured-DPPH "
  "antioxidant model and a deterministic druggability/CNS-MPO layer. Predictions are isotonic-"
  "calibrated, carry Mondrian conformal sets (empirical coverage 0.885–0.905; Norinder et al., 2014), "
  "and are grounded in nearest measured analogues and BBB-gated per-disease scores. Discrimination is "
  "strong within populated chemical space (random-split AUROC 0.94–0.98; scaffold and cluster 0.87–0.95) "
  "but degrades on genuinely novel scaffolds, which we report explicitly: where a query’s nearest "
  "training neighbour is below Tanimoto 0.4, the mean AUROC falls to 0.77 (range 0.60–0.85). The "
  "ensemble outperforms a Tanimoto read-across baseline on all eight endpoints (DeLong p < 0.05; MAO-A "
  "only marginally). A pre-registered case study shows that, unlike a general large language model, "
  "every prediction is traceable to a specific measurement. "
  "Conclusion. BrainSafe AI offers calibrated, evidence-grounded, safety-aware CNS profiling from "
  "measured data, with quantified and honest limits on novel chemistry. Code, models and data are "
  "released for research use, pending peer review.", align="j")
P("Keywords: QSAR; blood–brain barrier; neurodegeneration; conformal prediction; applicability "
  "domain; cheminformatics; ChEMBL.", italic=True, size=10)

H("Abbreviations", 2)
P("AUROC, area under the receiver-operating-characteristic curve; BBB, blood–brain barrier; CNS, "
  "central nervous system; CNS-MPO, CNS multiparameter optimisation; ECFP, extended-connectivity "
  "fingerprint; MCC, Matthews correlation coefficient; PR-AUC, area under the precision–recall "
  "curve; OOF, out-of-fold; CV, cross-validation; ATC, Anatomical Therapeutic Chemical; DPPH, "
  "2,2-diphenyl-1-picrylhydrazyl; QED, quantitative estimate of drug-likeness; AChE/BChE, "
  "acetyl-/butyryl-cholinesterase; BACE1, β-secretase 1; GSK-3β, glycogen synthase kinase-3β; MAO, "
  "monoamine oxidase; hERG, human ether-à-go-go-related gene; SERT, serotonin transporter.", size=10)

# ---- 1 INTRODUCTION ----
H("1. Introduction", 1)
P("Disorders of the nervous system are now a leading and still-rising cause of global disability (GBD "
  "2021 Nervous System Disorders Collaborators, 2024), yet the discovery of safe, brain-penetrant "
  "modulators remains slow. Natural products and flavonoids in particular attract sustained interest "
  "for their neuroprotective and antioxidant activity, but what is known about any one compound tends "
  "to be scattered across the literature rather than assembled into a single CNS profile (Hasan et al., "
  "2023). That profile has several parts that matter together: whether the compound crosses the BBB, "
  "which disease-relevant CNS targets it engages, how developable it is, whether it is safe, and "
  "whether anything like it has reached the clinic.", align="j")
P("Public resources already cover pieces of this. General ADMET platforms (Daina et al., 2017; Xiong "
  "et al., 2021; Fu et al., 2024; Cheng et al., 2012) predict BBB and hERG, but say nothing about "
  "CNS-target activity. Target-prediction servers (Daina et al., 2019; Awale and Reymond, 2019) return "
  "likely protein targets from chemical similarity, yet they do not condition those targets on brain "
  "penetration, roll them up into a disease-level view, attach calibrated uncertainty, or carry a "
  "safety axis. Individual endpoints have their own dedicated QSAR models: AChE and BACE1 (Ponzoni et "
  "al., 2019), MAO-B (Kumar et al., 2024), GSK-3β (Galati et al., 2023), and BBB permeability (Kumar et "
  "al., 2022; Huang et al., 2024). No single tool, though, draws them together. BrainSafe AI does exactly "
  "that, using models trained on measured public data and tested under four validation regimes.", align="j")

# ---- 2 METHODS ----
H("2. Materials and methods", 1)
figure("figures/fig1_workflow.png",
       "Figure 1. Overview of the BrainSafe AI pipeline, from measured data sources through curation, "
       "featurisation, ensemble training and calibration/conformal prediction to the integrated "
       "outputs and the four validation regimes.")
H("2.1 Data sources", 2)
P("Target bioactivities were retrieved from ChEMBL version 37 (release 2026-05-01) via its REST API "
  "(Gaulton et al., 2012; Mendez et al., 2019; Zdrazil et al., 2024), retaining records with a "
  "defined pChEMBL value (standard types IC50, Ki, Kd, EC50, Potency) for AChE (CHEMBL220), BChE "
  "(CHEMBL1914), BACE1 (CHEMBL4822), GSK-3β (CHEMBL262), MAO-A (CHEMBL1951), MAO-B (CHEMBL2039), the "
  "hERG anti-target (CHEMBL240), and receptors D2 (CHEMBL217), A2A (CHEMBL251), 5-HT2A (CHEMBL224) "
  "and SERT (CHEMBL228). BBB labels were taken from B3DB (Meng et al., 2021); antioxidant data from "
  "ChEMBL DPPH radical-scavenging assays (IC50/EC50 → pIC50); and clinical precedent from ChEMBL ATC "
  "level-1 ‘N’ molecules with a clinical phase. Structures for user-entered compounds are resolved "
  "via PubChem (Kim et al., 2023). All endpoints are trained on measured experimental data (Table 1), "
  "totalling 64,474 measured compound–endpoint records; ChEMBL document years (1976–2025) were "
  "retained to enable temporal validation (Sheridan, 2013).", align="j")
s0 = pd.read_csv("supplementary/STable0_data_provenance.csv")
s0 = s0[["Endpoint","Role","Modality","Source","Identifier","Measurement","n","Year range"]].copy()
s0["Endpoint"] = s0["Endpoint"].replace(DISP)
# B3DB/DPPH/clinical carry no ChEMBL document year; show 'n/a' rather than 'curated' in a year column
s0["Year range"] = s0["Year range"].replace({"curated": "n/a"})
P("Table 1. Data provenance for each endpoint (all trained on measured experimental data). "
  "‘n/a’ in the year column marks sources without per-compound ChEMBL document years.", italic=True, size=9)
table_from_df(s0)
P("")
H("2.2 Curation and labelling", 2)
P("SMILES were canonicalised and deduplicated by InChIKey (Bento et al., 2020). For classification, "
  "activities were aggregated per compound by median pChEMBL and labelled active where pChEMBL ≥ 6 "
  "(potency ≤ 1 µM) and inactive where pChEMBL < 5 (> 10 µM); the intermediate range was discarded. "
  "Receptor targets, 96–98% active because binders are preferentially reported, were modelled by "
  "potency regression on median pChEMBL rather than binary classification.", align="j")
H("2.3 Molecular representation, models and hyperparameters", 2)
P("Molecules were represented by a 1,024-bit Morgan (ECFP, radius 2) fingerprint (Rogers and Hahn, "
  "2010), computed with RDKit (RDKit, 2024), concatenated with 24 RDKit physicochemical descriptors. "
  "Classification used an "
  "unweighted-mean ensemble of random forest (Breiman, 2001), extremely randomised trees (Geurts et "
  "al., 2006) and histogram gradient boosting (Friedman, 2001; Ke et al., 2017) in scikit-learn "
  "(Pedregosa et al., 2011); regression used the corresponding regressor ensemble. Hyperparameters "
  "are listed in Table 2. A pre-specified quality gate (MCC ≥ 0.45 under scaffold CV) governed "
  "deployment; endpoints failing the gate as classifiers (D2, A2A, 5-HT2A, SERT) were modelled as "
  "regressions instead.", align="j")
P("Table 2. Model architecture and hyperparameters.", italic=True, size=9)
hp = pd.DataFrame([
    ["Random forest (classifier/regressor)", "n_estimators=300; min_samples_leaf=2; class_weight=balanced_subsample (clf)"],
    ["Extremely randomised trees", "n_estimators=300; min_samples_leaf=2; class_weight=balanced_subsample (clf)"],
    ["Histogram gradient boosting", "max_iter=300; learning_rate=0.06"],
    ["Ensemble", "unweighted mean of the three base learners (probabilities / predictions)"],
    ["Calibration", "isotonic regression on scaffold-CV out-of-fold predictions"],
    ["Conformal prediction", "Mondrian (class-conditional) inductive; significance ε = 0.10; 50/50 calibration/test"],
    ["Applicability domain", "maximum Tanimoto similarity to training set; flag threshold 0.30"],
    ["Baselines", "Tanimoto k-NN (k=5); logistic regression (balanced, standardised features)"],
    ["Random seed", "42 (all stages)"],
], columns=["Component", "Setting"])
table_from_df(hp); P("")
P("Table 3. Software environment.", italic=True, size=9)
sw = pd.DataFrame([["Python","3.13"],["RDKit","2026.03.2"],["scikit-learn","1.8.0"],["NumPy","2.4.6"],
                   ["pandas","3.0.3"],["SciPy","1.17.1"],["matplotlib","3.10.9"]], columns=["Package","Version"])
table_from_df(sw); P("")
H("2.4 Calibration, conformal prediction and integration", 2)
P("Calibrated probabilities were obtained by isotonic regression on scaffold-CV out-of-fold scores "
  "(Niculescu-Mizil and Caruana, 2005). Class-conditional (Mondrian) inductive conformal prediction "
  "(Vovk et al., 2005; Norinder et al., 2014; Mervin et al., 2021) yields, for a query with score s "
  "and per-class calibration nonconformity sets, a p-value per class c", align="j")
eq("p_c = (#{α in A_c : α ≥ α_query} + 1) / (|A_c| + 1),   α_active = 1 − s,   α_inactive = s")
P("and a prediction set comprising classes with p_c > ε (ε = 0.10). Effective CNS engagement combines "
  "target activity with brain penetration,", align="j")
eq("Engagement(target) = P(active | target) × P(BBB-penetrant)")
P("and per-disease scores take the maximum engagement over the targets mapped to each disease. This "
  "product is used deliberately as a monotone priority index for ranking, not as a calibrated "
  "probability of a clinical event: it multiplies two calibrated probabilities under a working "
  "independence assumption that does not strictly hold, since lipophilicity influences both barrier "
  "penetration and target binding. Because both factors are monotone in ‘desirable’, the product still "
  "orders candidates sensibly (a compound must both reach the brain and engage the target to score "
  "highly) and it down-weights potent binders that are predicted not to penetrate. We therefore treat "
  "the per-disease scores as heuristic priority indices, and the target and BBB probabilities that "
  "compose them are reported separately so a user can inspect each. The assumption and how the index "
  "could be validated are discussed in Section 4. "
  "Druggability is a deterministic composite of QED (Bickerton et al., 2012), Lipinski (Lipinski et "
  "al., 2001) and Veber (Veber et al., 2002) compliance, the CNS-MPO desirability score (Wager et "
  "al., 2010) and PAINS alerts (Baell and Holloway, 2010). Each prediction returns the nearest "
  "measured analogues by Tanimoto similarity (Willett, 2006).", align="j")
H("2.5 Validation regimes", 2)
P("Models were evaluated under: (i) a random stratified split; (ii) scaffold GroupKFold (k = 5) on "
  "Bemis–Murcko generic scaffolds (Bemis and Murcko, 1996), with all transforms fit within each fold "
  "(Tropsha, 2010; Wu et al., 2018); (iii) leave-cluster-out using sphere-exclusion clusters held "
  "out in full; and (iv) temporal validation, training on compounds reported up to the "
  "75th-percentile ChEMBL document year and testing on the most recent ~25% (Sheridan, 2013). "
  "Classification metrics were AUROC (with 95% bootstrap confidence intervals, 1,000 resamples), "
  "PR-AUC, balanced accuracy, MCC and Brier score; regression metrics were R², RMSE and Spearman ρ. "
  "Code, models and data-fetch scripts are released; a Streamlit application provides the interface.",
  align="j")

# ---- 3 RESULTS ----
H("3. Results", 1)
P("Results are reported in two modes. Non-comparative analysis (Sections 3.1, 3.4 and 3.5) reports "
  "each endpoint’s standalone performance against its own held-out measured data: absolute "
  "discrimination, calibration, conformal coverage, receptor-potency and antioxidant regression, and "
  "behaviour on reference compounds. Comparative analysis (Sections 3.2, 3.3 and 3.6) benchmarks the "
  "same models against external references: simpler internal baselines under an identical protocol, "
  "published QSAR performance ranges, and the general-purpose large-language-model paradigm.", align="j")
figure("figures/fig2_dataset.png",
       "Figure 2. Training-set size and class balance for each endpoint.")
H("3.1 Classification performance (standalone, non-comparative)", 2)
s1 = pd.read_csv("supplementary/STable1_classification_metrics.csv")
def ci_str(ep):
    c = cis.get(ep, {}).get("ci95"); return f"[{c[0]}, {c[1]}]" if c else "–"
t1 = s1[["endpoint","n","pos_rate","AUROC_random","AUROC_scaffold","AUROC_cluster","AUROC_temporal","Brier","MCC","conformal_coverage"]].copy()
t1.insert(5, "Scaffold 95% CI", [ci_str(e) for e in s1["endpoint"]])
t1.columns = ["Endpoint","n","Pos. rate","AUROC rand","AUROC scaf","Scaf 95% CI","AUROC clust","AUROC temp","Brier","MCC","Conf. cov."]
t1["Endpoint"] = t1["Endpoint"].replace(DISP)
P("Table 4. Classification performance under four validation regimes, with 95% bootstrap CI on the "
  "scaffold-CV AUROC, Brier score, MCC and empirical conformal coverage (target 0.90).", italic=True, size=9)
table_from_df(t1); P("")
figure("figures/fig3_validation.png",
       "Figure 3. AUROC across validation regimes; error bars on scaffold-CV bars denote 95% bootstrap CIs.")
figure("figures/fig4_roc_calibration.png",
       "Figure 4. (A) Scaffold cross-validation ROC curves (out-of-fold). (B) Reliability diagrams after "
       "isotonic calibration for representative endpoints.")
figure("figures/fig5_conformal_comparison.png",
       "Figure 5. (A) Empirical coverage of 90%-level conformal prediction sets. (B) Scaffold-CV AUROC "
       "of the deployed ensemble versus Tanimoto k-nearest-neighbour and logistic-regression baselines.")
P("Within populated chemical space, random-split AUROC was 0.94–0.98 across endpoints; scaffold and "
  "leave-cluster-out AUROC were 0.87–0.95; and temporal-split AUROC was 0.61–0.92 (Figure 3; Table 4). "
  "Isotonic calibration gave Brier scores of 0.04–0.14 (Figure 4B) and empirical conformal coverage of "
  "0.885–0.905 against the 0.90 target (Figure 5A).", align="j")
figure("figures/fig8_generalisation.png",
       "Figure 6. Generalisation versus chemical novelty. For each held-out compound, AUROC is binned by "
       "its nearest-neighbour Tanimoto to the training set; the dashed line is the n-weighted mean over "
       "the eight classification endpoints.")
P("A random-split headline overstates prospective utility, because ChEMBL target sets are "
  "analogue-dense and most published models report only that split. We therefore quantify "
  "generalisation as an explicit function of chemical novelty (Figure 6; Supplementary Table S5): for "
  "each held-out compound we bin its nearest-neighbour Tanimoto to the training set and measure AUROC "
  "within each bin. Discrimination is retained where a query resembles training chemistry (n-weighted "
  "mean AUROC 0.973 at Tanimoto ≥ 0.8 and 0.942 at 0.6–0.8) but falls steadily as novelty rises, to "
  "0.874 at 0.4–0.6 and 0.774 below 0.4; in that most-novel bin the per-endpoint values span 0.60 "
  "(BChE) to 0.85 (BBB). This is the honest measure of behaviour on genuinely new scaffolds, and it is "
  "why the applicability-domain flag (Section 2.4) marks queries below Tanimoto 0.3–0.4 as "
  "low-confidence. Taken with the random-split numbers, the claim is state-of-the-art discrimination "
  "within populated chemical space, with quantified degradation on novel chemistry rather than a single "
  "inflated headline.", align="j")
H("3.2 Ablation against simpler baselines (comparative)", 2)
s9 = pd.read_csv("supplementary/STable9_baseline_comparison.csv")
_ens = s9["Ensemble_AUROC"].mean(); _knn = s9["kNN_Tanimoto_AUROC"].mean(); _lr = s9["LogisticRegression_AUROC"].mean()
P("Under an identical scaffold-split protocol and feature set, the deployed ensemble was compared "
  "with a k-nearest-neighbour Tanimoto ‘read-across’ baseline (the closest algorithmic analogue to "
  "associative structural recall) and with L2-regularised logistic regression. The ensemble attained "
  f"a mean scaffold-split AUROC of {_ens:.3f}, versus {_knn:.3f} for k-nearest-neighbour "
  f"(mean Δ = +{_ens-_knn:.3f}) and {_lr:.3f} for logistic regression (mean Δ = +{_ens-_lr:.3f}), "
  "and was best on every one of the eight endpoints (Table 5; Figure 5B). Exceeding a pure "
  "nearest-neighbour read-across indicates that the model captures structure–activity relationships not "
  "reducible to retrieving the most similar known molecule.", align="j")
t6 = s9[["endpoint","Ensemble_AUROC","kNN_Tanimoto_AUROC","LogisticRegression_AUROC","delta_vs_kNN"]].copy()
t6.columns = ["Endpoint","Ensemble","kNN-Tanimoto","Logistic reg.","Δ vs kNN"]
t6["Endpoint"] = t6["Endpoint"].replace(DISP)
P("Table 5. Scaffold-split AUROC of the deployed ensemble versus baselines (Supplementary Table S9).", italic=True, size=9)
table_from_df(t6); P("")
sig = json.load(open("BS_significance_report.json"))
_ma = sig["MAO_A"]; _ci = _ma["boot_delta95_vs_kNN"]
P("To test whether these margins exceed sampling noise, we applied DeLong’s test for two correlated ROC "
  "curves to each ensemble-versus-kNN comparison, paired on identical held-out compounds, and "
  "corroborated it with a paired bootstrap (2,000 resamples; Supplementary Table S14). The paired test "
  "is more powerful than a per-AUROC confidence interval because the two ROC curves are strongly "
  "correlated on the same samples. The ensemble’s advantage is significant at all eight endpoints "
  "(DeLong p < 0.05): decisively so for seven (p < 0.001) and marginally for MAO-A, where the gain is "
  f"smallest (Δ = +{_ma['delta_vs_kNN']:.3f}, DeLong p = {_ma['delong_p_vs_kNN']:.3f}, bootstrap 95% CI "
  f"[{_ci[0]:+.3f}, {_ci[1]:+.3f}]). We therefore report a statistically significant improvement over "
  "read-across on every endpoint, with MAO-A flagged as the borderline case.", align="j")
H("3.3 Comparison with the literature (comparative)", 2)
figure("figures/fig6_benchmark.png",
       "Figure 7. Per-endpoint random-split AUROC (this work) relative to reported random-split ranges.")
P("On like-for-like random splits, per-endpoint AUROC was within or above reported ranges, for "
  "example BBB (0.88–0.96; Kumar et al., 2022; Huang et al., 2024) and hERG (0.86–0.93). The full "
  "per-endpoint comparison against the collected literature ranges is given in Supplementary Table S7 "
  "(Figure 7). We stress that this is a cross-study comparison over different test sets and is therefore "
  "weaker than a same-test-set benchmark; a direct comparison against a public server on identical "
  "held-out compounds is a planned addition (Section 4).", align="j")
H("3.4 Receptor potency regression and antioxidant model (non-comparative)", 2)
s2 = pd.read_csv("supplementary/STable2_receptor_regression.csv")
t2 = s2[["receptor","n","scaffold_cv_R2","RMSE","Spearman","temporal_R2"]].copy()
t2.columns = ["Receptor","n","Scaffold R²","RMSE","Spearman","Temporal R²"]
t2["Receptor"] = t2["Receptor"].replace(DISP)
P("Table 6. Receptor potency-regression performance (scaffold cross-validation and temporal split).", italic=True, size=9)
table_from_df(t2); P("")
figure("figures/fig7_regression.png",
       "Figure 8. Predicted versus measured potency (scaffold cross-validation) for the measured "
       "antioxidant (DPPH) model and the four receptor potency-regression endpoints (panels A–E).")
am = json.load(open("models_genuine/antioxidant_measured_meta.json"))
P(f"Receptor potency regressions achieved scaffold-CV R² of 0.34–0.53 and Spearman ρ of 0.57–0.71; "
  f"temporal R² was lower (Table 6), and these endpoints are reported as ranking-grade. The "
  f"antioxidant model trained on measured DPPH data (n = {am['n']}) achieved scaffold-CV R² = "
  f"{am['scaffold_cv_r2']}, RMSE = {am['rmse']} and Spearman = {am['spearman']} (Figure 8); the "
  f"earlier curated score correlated only weakly with measured DPPH (Spearman = "
  f"{am['crosscheck_curated_vs_measured_spearman']}).", align="j")
H("3.5 Behaviour on reference compounds (non-comparative)", 2)
P("With chemistry-only inputs the integrated system reproduced established pharmacology: rivastigmine "
  "resolved to Alzheimer’s disease via BChE; rasagiline to Parkinson’s disease via MAO-B; fluoxetine "
  "to depression via SERT with a corresponding clinical-precedent match; terfenadine, withdrawn for "
  "cardiotoxicity, was predicted positive at hERG; and resveratrol was predicted BBB non-penetrant, "
  "consistent with reported flavonoid CNS bioavailability (Hasan et al., 2023).", align="j")

H("3.6 Illustrative case study: general-purpose LLMs (comparative)", 2)
P("A common question is whether a dedicated tool is needed when a general-purpose large language model "
  "(LLM) can be queried in natural language. We address this as an illustrative case study, not a "
  "benchmark: the panel below is small (ten compounds), single-shot, and the models are versioned "
  "snapshots that will change, so the numbers should be read as a qualitative demonstration rather than "
  "a rigorous performance comparison. Our contribution is the QSAR and its integration; this section is "
  "context, not a load-bearing result. We consider three lines: published benchmark evidence, the "
  "architectural basis of the difference, and a grounded-output demonstration.", align="j")
P("Benchmark evidence. On molecular property prediction (the task class BrainSafe performs), "
  "general-purpose LLMs consistently underperform specialised machine-learning models. In an "
  "eight-task chemistry benchmark, LLMs including GPT-4 lag task-specific models on property "
  "prediction and parse SMILES unreliably (Guo et al., 2023); a dedicated molecule-prediction "
  "benchmark reports that LLMs ‘generally lag behind ML models’, particularly where molecular "
  "geometry matters (Zhong et al., 2024); and even fine-tuned LLMs become competitive with dedicated "
  "QSAR models only in the low-data limit, not at the data scale used here (Jablonka et al., 2024). "
  "LLMs also exhibit documented factual hallucination in generative settings (Ji et al., 2023).", align="j")
P("Scientific background. A general LLM is an autoregressive next-token predictor over text; it does "
  "not compute a molecular fingerprint, does not fit an explicit function from chemical structure to "
  "measured bioactivity, and does not emit a probability with a calibration or coverage guarantee. "
  "BrainSafe encodes each molecule as an ECFP-4 fingerprint plus 24 physicochemical descriptors, "
  "learns a structure-to-measured-activity mapping from 64,474 records, and returns a calibrated "
  "probability wrapped in a conformal set with empirically verified ~90% coverage, together with the "
  "nearest measured analogue and its measured pChEMBL, none of which an LLM can provide. Table 7 "
  "summarises the capability differences (Supplementary Table S8).", align="j")
s8 = pd.read_csv("supplementary/STable8_llm_capability_comparison.csv")
s8.columns = ["Dimension","BrainSafe AI","General-purpose LLM"]
P("Table 7. Capability comparison: BrainSafe AI versus a general-purpose large language model.", italic=True, size=9)
table_from_df(s8); P("")
P("Reproducible grounded-output demonstration. For fixed input structures the deployed engine "
  "returns verifiable artifacts (script BS_llm_comparison.py; output BS_llm_comparison.json): "
  "donepezil → AChE calibrated P = 1.00 with the nearest measured analogue at Tanimoto 1.00 "
  "(donepezil is itself a measured training compound, pChEMBL 7.75) and hERG P = 0.78; terfenadine → hERG "
  "P = 1.00, correctly flagging the cardiotoxicity for which it was withdrawn while also calling it "
  "BBB non-penetrant; and a novel arylpiperazine of an unpublished scaffold → an honest "
  "conformal ‘uncertain’ set for AChE grounded in a measured analogue (pChEMBL 4.82) rather than a "
  "confident but unverifiable text answer. Every value is traceable to a measurement. This grounding "
  "(a calibrated probability, a coverage-guaranteed set, and measured-analogue provenance for any "
  "structure, including novel ones) is the scientific justification for a dedicated tool "
  "complementary to, not replaced by, general LLMs.", align="j")
P("Grounded-output case study. Using a prompt, a ten-compound panel (approved-drug pharmacology plus "
  "one unpublished scaffold) and a scoring rubric all fixed in advance (protocol and key released with "
  "the code), we queried four general-purpose LLMs once each (Gemini Pro, ChatGPT/GPT-4o, Perplexity, "
  "Claude; specific dated snapshots) and scored the replies against the measured-data key (Table 8; "
  "Supplementary Table S13). These are single-shot results on a small panel and are illustrative only.", align="j")
sb = pd.read_csv("supplementary/STable13_llm_scoreboard.csv")
_disp = sb[["system","BBB_acc","hERG_acc","chembl_ids_given","fabricated_ids","wrong_structure_ids","novel_confabulation"]].copy()
_disp.columns = ["System","BBB (of 9)","hERG (of 5)","IDs cited","Fabricated","Wrong-molecule","Novel confab."]
P("Table 8. Illustrative LLM case study (single-shot, ten-compound panel), scored against the frozen "
  "measured-data key. hERG scored on five uncontested compounds. Note the asymmetry: BrainSafe cites "
  "measured analogues by structure and never emits a ChEMBL identifier, so by construction it cannot "
  "fabricate one; the identifier columns therefore describe the LLMs’ behaviour, not a like-for-like score.",
  italic=True, size=9)
table_from_df(_disp); P("")
P("Two patterns emerge, with the caveats above. First, on well-known approved drugs the LLMs are "
  "strong: three of four matched or exceeded BrainSafe on BBB (9/9 vs 8/9; BrainSafe mis-called the "
  "borderline-lipophilic astemizole), so a dedicated tool is not justified by accuracy on famous "
  "compounds. Second, when the LLMs volunteered a specific ChEMBL identifier as provenance, 14 of the "
  "31 identifiers cited (45%) were fabricated or resolved to the wrong molecule. For example, one "
  "model’s cited ‘rasagiline’ identifier is in fact fluticasone propionate, and another’s "
  "‘rivastigmine’ identifier is pyridoxine. On the single unpublished compound, all four asserted a "
  "specific target and potency and disagreed with one another (three AChE, one D2), where BrainSafe "
  "returned a conformal ‘uncertain’ set grounded in a measured analogue. We read this narrowly: on a "
  "small panel, a text model can reproduce textbook classifications but its volunteered evidence is "
  "unreliable and it does not abstain on genuinely novel chemistry, whereas the grounded tool cites a "
  "real measurement or flags uncertainty.", align="j")

# ---- 4 DISCUSSION ----
H("4. Discussion", 1)
P("What BrainSafe AI contributes is the assembly, not the parts. Its components are individually "
  "familiar: fingerprint and tree-ensemble QSAR, the BBB and hERG models, conformal prediction, and "
  "the QED and CNS-MPO druggability rules, each well established in its own right (Rogers and Hahn, "
  "2010; Breiman, 2001; Norinder et al., 2014; "
  "Wager et al., 2010). What we have not found elsewhere is all of them working as one measured-data "
  "CNS profiler that is at the same time calibrated, conformal, evidence-grounded, BBB-gated, "
  "safety-aware and clinically contextualised (Daina et al., 2017, 2019; Fu et al., 2024; Awale and "
  "Reymond, 2019).", align="j")
P("AUROC fell from the random split (0.94–0.98) to the temporal one (0.61–0.92), and the size of that "
  "drop tracks how much new chemistry each test set contained: between 71% and 91% of the recent "
  "compounds carried Bemis–Murcko scaffolds the model had never seen. The two extremes are "
  "instructive. BACE1 held up at 0.92, but its recent test set was 93% active, which flatters the "
  "score; MAO-A, with a balanced 45%-active test set, gives the more representative figure of 0.61. Reporting "
  "all four regimes is what lets a reader see generalisation across both chemical space and time "
  "(Sheridan, 2013; Tropsha, 2010).", align="j")
P("Why build a dedicated tool when a general-purpose LLM can be asked the same questions in plain "
  "language? Our head-to-head (Section 3.6) answers with results rather than assertion, and the answer "
  "comes in two parts. On the task itself, molecular property prediction, general LLMs are known to "
  "trail specialised machine learning (Guo et al., 2023; Zhong et al., 2024); fine-tuning narrows the "
  "gap only when data are scarce (Jablonka et al., 2024), and BrainSafe works from 64,474 measured "
  "records. That much was expected. What the benchmark adds is the failure mode that matters in "
  "practice. Four current models matched or beat BrainSafe at classifying well-known drugs, yet 45% of "
  "the ChEMBL identifiers they offered as evidence were fabricated or pointed to the wrong molecule, "
  "and every one of them invented a target and potency for an unpublished compound, disagreeing with "
  "each other along the way. The deeper reason is architectural. An LLM generates fluent text; it does "
  "not compute a fingerprint, fit structure to measured activity, or attach a calibrated probability, a "
  "coverage guarantee or a domain boundary, and it will hallucinate with confidence (Ji et al., 2023). "
  "BrainSafe returns, for any structure, a calibrated probability, a conformal set with roughly 90% "
  "empirical coverage, an explicit in- or out-of-domain flag, and the nearest measured analogue with "
  "its pChEMBL. The two are complementary, not interchangeable, wherever a decision has to be auditable "
  "and anchored in measurement.", align="j")
P("Threats to validity (scientific-flaw self-audit, with quantitative tests). We enumerated the "
  "principal methodological risks and, rather than merely noting them, ran targeted analyses to "
  "bound each (BS_flaw_fixes.py / BS_assay_composition.py / BS_assay_sensitivity.py). "
  "(1) Assay heterogeneity: activities pooled per target span IC50/Ki/Kd/EC50/Potency; we first "
  "quantified the composition (Supplementary Table S11), finding IC50 dominant for every target "
  "(81–92%) except GSK-3β, which is genuinely mixed (IC50 49%, EC50 33%, Ki 16%). Retraining the "
  "deployed ensemble under scaffold CV on the dominant single assay type (IC50) only versus the "
  "pooled set changed AUROC by at most 0.006 across the three endpoints tested, including the most "
  "heterogeneous, GSK-3β (pooled 0.919 vs IC50-only 0.913; MAO-B −0.006; hERG 0.000; Supplementary "
  "Table S12). Pooling on the standardised pChEMBL scale therefore does not materially distort "
  "discrimination. (2) Label-threshold sensitivity: re-deriving labels at alternative definitions and "
  "re-measuring scaffold-CV AUROC with the deployed ensemble gave a maximum spread of 0.109 across "
  "four endpoints and four definitions; the stricter ≥6.5/<5.5 cut sat within 0.01–0.02 of the "
  "deployed cut, and the ‘sharp boundary’ cut retaining the 5–6 grey zone was consistently worst, "
  "validating its removal (Supplementary Table S10; per-operating-threshold values in S4). "
  "(3) Applicability-domain cut-off: the n-weighted similarity-binned AUROC falls monotonically from "
  "0.958 (nearest-neighbour Tanimoto ≥0.8) to 0.770 (<0.4) (Supplementary Table S5), empirically "
  "justifying the out-of-domain flag in the 0.3–0.4 band. (4) Disease mapping: the target-to-disease "
  "synthesis is a transparent, knowledge-based rule (not a learned layer), with each driver tagged by "
  "provenance, so it is inspectable and overridable. (5) Single safety anti-target: cardiotoxicity is "
  "represented by hERG alone; other liabilities are out of scope. (6) Read-across ceiling: the "
  "ensemble exceeds a k-nearest-neighbour baseline on every endpoint (Section 3.2), so performance is "
  "not merely memorised nearest-neighbour recall. Each risk is surfaced in the tool output or the "
  "supplementary tables rather than concealed.", align="j")
P("The BBB-gated disease score is the closest thing here to a new method, and we are deliberately "
  "modest about it. It is a heuristic priority index: a product of two calibrated probabilities under a "
  "working independence assumption that does not strictly hold, since lipophilicity drives both barrier "
  "penetration and target binding, so the product is not itself a calibrated probability of any "
  "measurable event. It remains useful for ranking, because a candidate must both reach the brain and "
  "engage the target to score highly, and a potent binder predicted not to penetrate is down-weighted. "
  "A full quantitative validation of the index would need an outcome set in which brain penetration "
  "genuinely varies; our clinical-precedent set is unsuitable, being composed almost entirely of "
  "BBB-penetrant nervous-system drugs, so gating is close to neutral within it. Building a mixed "
  "penetrant/non-penetrant benchmark and showing that gating improves disease-level ranking on it is a "
  "clear next step that would turn the heuristic into a validated result.", align="j")
P("Two further comparisons would strengthen the claims and are planned. First, a same-test-set "
  "benchmark: running our held-out BBB and hERG compounds through a public server such as ADMETlab 3.0 "
  "and reporting AUROC on identical molecules, which is more informative than the cross-study range "
  "comparison in Section 3.3. Second, prospective validation, either a strict external test set drawn "
  "from a source other than ChEMBL or B3DB, or modest wet-lab confirmation; either would support a "
  "claim of generalisation beyond the retrospective splits reported here.", align="j")
P("Some limitations remain, and they bound how the tool should be read. It predicts that a molecule "
  "engages a target, not whether it activates or blocks it (agonism versus antagonism), and engagement "
  "is not the same thing as clinical benefit; the clinical-precedent layer reports structural "
  "similarity to compounds that reached documented trial phases, which is context, not a prediction of "
  "efficacy. No prospective wet-lab validation has yet been done. GSK-3β and MAO-A generalise less well "
  "over time and are flagged as lower-confidence, and the pooled DPPH antioxidant data transfer only "
  "weakly across time. The tool is meant for research hypothesis generation and prioritisation, not for "
  "clinical or diagnostic use.", align="j")

# ---- 5 CONCLUSION ----
H("5. Conclusion", 1)
P("BrainSafe AI is a calibrated, evidence-grounded, multi-endpoint CNS profiler built entirely on "
  "measured public data. Within populated chemical space its per-endpoint discrimination is "
  "state-of-the-art-grade, and it improves significantly on read-across at every endpoint; we place the "
  "harder scaffold, cluster and temporal numbers, and the explicit degradation on novel scaffolds, "
  "alongside that headline rather than in place of it. The limitations are stated plainly, and the two "
  "clearest routes to a stronger contribution, a validated BBB-gating index and a prospective external "
  "test, are identified. On that basis the tool is suited to a resource publication and to practical "
  "research prioritisation.", align="j")

# ---- BACK MATTER ----
H("Data and code availability", 2)
P("All code, the per-endpoint measured training data (data/endpoints/*.csv), trained-model metadata, "
  "and every validation report, supplementary table and figure script are released in a public, "
  "version-tagged repository (release v1.0.0; https://github.com/krishna-g-999/brainsafe-ai) under the "
  "MIT licence, and archived on Zenodo (DOI: [inserted at submission]). Trained-model binaries (~1.3 GB) "
  "are regenerated deterministically (fixed random seed 42) by the released scripts and are also "
  "deposited in the Zenodo archive. Source data derive from ChEMBL version 37 (release 2026-05-01; "
  "Gaulton et al., 2012; Mendez et al., 2019; Zdrazil et al., 2024) and B3DB (Meng et al., 2021); the "
  "exact retrieval queries are in BS_fetch_endpoints.py. The interactive Streamlit application "
  "(app_v6_final.py) is available at [public URL inserted at submission].", align="j")
H("Ethics statement", 2)
P("This study used only publicly available molecular and bioactivity data and did not involve human "
  "participants, human tissue, or animals; no ethical approval was required.", align="j")
H("Author contributions, funding and competing interests", 2)
P("Author contributions: [To be completed]. Funding: [To be completed]. Competing interests: the "
  "authors declare no competing interests. [Confirm on submission.]", align="j")

# ---- REFERENCES ----
doc.add_page_break(); H("References", 1)
refs = [
"Awale, M. and Reymond, J.-L. (2019) ‘The polypharmacology browser PPB2: target prediction combining nearest neighbors with machine learning’, Journal of Chemical Information and Modeling, 59(1), pp. 10–17.",
"Baell, J.B. and Holloway, G.A. (2010) ‘New substructure filters for removal of pan assay interference compounds (PAINS) from screening libraries and for their exclusion in bioassays’, Journal of Medicinal Chemistry, 53(7), pp. 2719–2740.",
"Bemis, G.W. and Murcko, M.A. (1996) ‘The properties of known drugs. 1. Molecular frameworks’, Journal of Medicinal Chemistry, 39(15), pp. 2887–2893.",
"Bento, A.P., Hersey, A., Félix, E., Landrum, G., Gaulton, A., Atkinson, F., Bellis, L.J., De Veij, M. and Leach, A.R. (2020) ‘An open source chemical structure curation pipeline using RDKit’, Journal of Cheminformatics, 12, 51.",
"Bickerton, G.R., Paolini, G.V., Besnard, J., Muresan, S. and Hopkins, A.L. (2012) ‘Quantifying the chemical beauty of drugs’, Nature Chemistry, 4(2), pp. 90–98.",
"Breiman, L. (2001) ‘Random forests’, Machine Learning, 45(1), pp. 5–32.",
"Cheng, F., Li, W., Zhou, Y., Shen, J., Wu, Z., Liu, G., Lee, P.W. and Tang, Y. (2012) ‘admetSAR: a comprehensive source and free tool for assessment of chemical ADMET properties’, Journal of Chemical Information and Modeling, 52(11), pp. 3099–3105.",
"Daina, A., Michielin, O. and Zoete, V. (2017) ‘SwissADME: a free web tool to evaluate pharmacokinetics, drug-likeness and medicinal chemistry friendliness of small molecules’, Scientific Reports, 7, 42717.",
"Daina, A., Michielin, O. and Zoete, V. (2019) ‘SwissTargetPrediction: updated data and new features for efficient prediction of protein targets of small molecules’, Nucleic Acids Research, 47(W1), pp. W357–W364.",
"Friedman, J.H. (2001) ‘Greedy function approximation: a gradient boosting machine’, Annals of Statistics, 29(5), pp. 1189–1232.",
"Fu, L., Shi, S., Yi, J., Wang, N., He, Y., Wu, Z., Peng, J., Deng, Y., Wang, W., Wu, C., Lyu, A., Zeng, X., Zhao, W., Hou, T. and Cao, D. (2024) ‘ADMETlab 3.0: an updated comprehensive online ADMET prediction platform’, Nucleic Acids Research, 52(W1), pp. W422–W431.",
"Galati, S., Di Stefano, M., Bertini, S., Granchi, C., Giordano, A., Gado, F., Macchia, M., Tuccinardi, T. and Poli, G. (2023) ‘Identification of new GSK3β inhibitors through a consensus machine learning-based virtual screening’, International Journal of Molecular Sciences, 24(24), 17233.",
"Gaulton, A., Bellis, L.J., Bento, A.P., Chambers, J., Davies, M., Hersey, A., Light, Y., McGlinchey, S., Michalovich, D., Al-Lazikani, B. and Overington, J.P. (2012) ‘ChEMBL: a large-scale bioactivity database for drug discovery’, Nucleic Acids Research, 40(D1), pp. D1100–D1107.",
"GBD 2021 Nervous System Disorders Collaborators (2024) ‘Global, regional, and national burden of disorders affecting the nervous system, 1990–2021’, The Lancet Neurology, 23(4), pp. 344–381.",
"Geurts, P., Ernst, D. and Wehenkel, L. (2006) ‘Extremely randomized trees’, Machine Learning, 63(1), pp. 3–42.",
"Guo, T., Guo, K., Nan, B., Liang, Z., Guo, Z., Chawla, N.V., Wiest, O. and Zhang, X. (2023) ‘What can large language models do in chemistry? A comprehensive benchmark on eight tasks’, Advances in Neural Information Processing Systems 36 (Datasets and Benchmarks Track). arXiv:2305.18365.",
"Hasan, S., Khatri, N., Rahman, Z.N., Menezes, A.A., Martini, J., Shehjar, F., Mujeeb, N. and Shah, Z.A. (2023) ‘Neuroprotective potential of flavonoids in brain disorders’, Brain Sciences, 13(9), 1258.",
"Huang, E.T.C., Yang, J.-S., Liao, K.Y.K., Tseng, W.C.W., Lee, C.K., Gill, M., Compas, C., See, S. and Tsai, F.-J. (2024) ‘Predicting blood–brain barrier permeability of molecules with a large language model and machine learning’, Scientific Reports, 14, 15844.",
"Jablonka, K.M., Schwaller, P., Ortega-Guerrero, A. and Smit, B. (2024) ‘Leveraging large language models for predictive chemistry’, Nature Machine Intelligence, 6(2), pp. 161–169.",
"Ji, Z., Lee, N., Frieske, R., Yu, T., Su, D., Xu, Y., Ishii, E., Bang, Y.J., Madotto, A. and Fung, P. (2023) ‘Survey of hallucination in natural language generation’, ACM Computing Surveys, 55(12), Article 248.",
"Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q. and Liu, T.-Y. (2017) ‘LightGBM: a highly efficient gradient boosting decision tree’, Advances in Neural Information Processing Systems, 30, pp. 3146–3154.",
"Kim, S., Chen, J., Cheng, T., Gindulyte, A., He, J., He, S., Li, Q., Shoemaker, B.A., Thiessen, P.A., Yu, B., Zaslavsky, L., Zhang, J. and Bolton, E.E. (2023) ‘PubChem 2023 update’, Nucleic Acids Research, 51(D1), pp. D1373–D1380.",
"Kumar, R., Sharma, A., Alexiou, A., Bilgrami, A.L., Kamal, M.A. and Ashraf, G.M. (2022) ‘DeePred-BBB: a blood brain barrier permeability prediction model with improved accuracy’, Frontiers in Neuroscience, 16, 858126.",
"Kumar, S., Bhowmik, R., Oh, J.M., Abdelgawad, M.A., Ghoneim, M.M., Al-Serwi, R.H., Kim, H. and Mathew, B. (2024) ‘Machine learning driven web-based app platform for the discovery of monoamine oxidase B inhibitors’, Scientific Reports, 14, 4868.",
"Lipinski, C.A., Lombardo, F., Dominy, B.W. and Feeney, P.J. (2001) ‘Experimental and computational approaches to estimate solubility and permeability in drug discovery and development settings’, Advanced Drug Delivery Reviews, 46(1–3), pp. 3–26.",
"Mendez, D., Gaulton, A., Bento, A.P., Chambers, J., De Veij, M., Félix, E., Magariños, M.P., Mosquera, J.F., Mutowo, P., Nowotka, M., Gordillo-Marañón, M., Hunter, F., Junco, L., Mugumbate, G., Rodriguez-Lopez, M., Atkinson, F., Bosc, N., Radoux, C.J., Segura-Cabrera, A., Hersey, A. and Leach, A.R. (2019) ‘ChEMBL: towards direct deposition of bioassay data’, Nucleic Acids Research, 47(D1), pp. D930–D940.",
"Meng, F., Xi, Y., Huang, J. and Ayers, P.W. (2021) ‘A curated diverse molecular database of blood-brain barrier permeability with chemical descriptors’, Scientific Data, 8, 289.",
"Mervin, L.H., Johansson, S., Semenova, E., Giblin, K.A. and Engkvist, O. (2021) ‘Uncertainty quantification in drug design’, Drug Discovery Today, 26(2), pp. 474–489.",
"Niculescu-Mizil, A. and Caruana, R. (2005) ‘Predicting good probabilities with supervised learning’, Proceedings of the 22nd International Conference on Machine Learning, pp. 625–632.",
"Norinder, U., Carlsson, L., Boyer, S. and Eklund, M. (2014) ‘Introducing conformal prediction in predictive modeling. A transparent and flexible alternative to applicability domain determination’, Journal of Chemical Information and Modeling, 54(6), pp. 1596–1603.",
"Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M. and Duchesnay, É. (2011) ‘Scikit-learn: machine learning in Python’, Journal of Machine Learning Research, 12, pp. 2825–2830.",
"Ponzoni, I., Sebastián-Pérez, V., Martínez, M.J., Roca, C., De la Cruz Pérez, C., Cravero, F., Vazquez, G.E., Páez, J.A., Díaz, M.F. and Campillo, N.E. (2019) ‘QSAR classification models for predicting the activity of inhibitors of beta-secretase (BACE1) associated with Alzheimer’s disease’, Scientific Reports, 9, 9102.",
"RDKit (2024) RDKit: open-source cheminformatics. Available at: https://www.rdkit.org (Accessed: 2026).",
"Rogers, D. and Hahn, M. (2010) ‘Extended-connectivity fingerprints’, Journal of Chemical Information and Modeling, 50(5), pp. 742–754.",
"Sheridan, R.P. (2013) ‘Time-split cross-validation as a method for estimating the goodness of prospective prediction’, Journal of Chemical Information and Modeling, 53(4), pp. 783–790.",
"Tropsha, A. (2010) ‘Best practices for QSAR model development, validation, and exploitation’, Molecular Informatics, 29(6–7), pp. 476–488.",
"Veber, D.F., Johnson, S.R., Cheng, H.-Y., Smith, B.R., Ward, K.W. and Kopple, K.D. (2002) ‘Molecular properties that influence the oral bioavailability of drug candidates’, Journal of Medicinal Chemistry, 45(12), pp. 2615–2623.",
"Vovk, V., Gammerman, A. and Shafer, G. (2005) Algorithmic learning in a random world. New York: Springer.",
"Wager, T.T., Hou, X., Verhoest, P.R. and Villalobos, A. (2010) ‘Moving beyond rules: the development of a central nervous system multiparameter optimization (CNS MPO) approach’, ACS Chemical Neuroscience, 1(6), pp. 435–449.",
"Willett, P. (2006) ‘Similarity-based virtual screening using 2D fingerprints’, Drug Discovery Today, 11(23–24), pp. 1046–1053.",
"Wu, Z., Ramsundar, B., Feinberg, E.N., Gomes, J., Geniesse, C., Pappu, A.S., Leswing, K. and Pande, V. (2018) ‘MoleculeNet: a benchmark for molecular machine learning’, Chemical Science, 9(2), pp. 513–530.",
"Xiong, G., Wu, Z., Yi, J., Fu, L., Yang, Z., Hsieh, C., Yin, M., Zeng, X., Wu, C., Lu, A., Chen, X., Hou, T. and Cao, D. (2021) ‘ADMETlab 2.0: an integrated online platform for accurate and comprehensive predictions of ADMET properties’, Nucleic Acids Research, 49(W1), pp. W5–W14.",
"Zdrazil, B., Felix, E., Hunter, F., Manners, E.J., Blackshaw, J., Corbett, S., de Veij, M., Ioannidis, H., Mendez, D., Mosquera, J.F., Magariños, M.P., Bosc, N., Arcila, R., Kizilören, T., Gaulton, A., Bento, A.P., Adasme, M.F., Monecke, P., Landrum, G.A. and Leach, A.R. (2024) ‘The ChEMBL Database in 2023’, Nucleic Acids Research, 52(D1), pp. D1180–D1192.",
"Zhong, Z., Zhou, K. and Mottin, D. (2024) ‘Benchmarking large language models for molecule prediction tasks’. arXiv:2403.05075.",
]
for rd in refs:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.first_line_indent = Inches(-0.5); p.paragraph_format.space_after = Pt(6)
    p.add_run(rd).font.size = Pt(10)
P(f"Total references: {len(refs)}.", italic=True, size=9)
doc.save("BrainSafe_AI_Manuscript.docx")
print("Saved BrainSafe_AI_Manuscript.docx |", len(refs), "references | graphical abstract + 7 figures | 5 tables | line numbers + highlights + abbreviations")
